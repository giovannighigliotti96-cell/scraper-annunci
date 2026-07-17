"""
Personalizzazione CV per offerte ad alto match (>=80%).

Attivato solo per le offerte che il match semantico di scraper.py ha già
valutato >=80%. Per queste:
  1. il title sotto il nome diventa il titolo esatto dell'annuncio
  2. il campo città nella sidebar (INFO E CONTACTS) diventa la città
     dell'annuncio, se è una delle città seguite (Genova/Milano/Torino)
  3. il paragrafo PROFILO PROFESSIONALE viene riformulato nello stesso
     framework (struttura di frase) ma calibrato sul job specifico
  4. i titoli di ruolo mostrati per ciascuna esperienza (l'etichetta del
     ruolo, non l'azienda) e i bullet delle esperienze vengono riformulati
     liberamente in base a quanto propone l'LLM

Vincoli non negoziabili (per scelta esplicita dell'utente, che fa una
revisione personale prima di ogni candidatura):
  - non si possono introdurre strumenti/tecnologie/sistemi nominati (es.
    Salesforce, SAP) che non compaiono già nel CV originale — l'unico
    controllo rigido della verifica
  - azienda, periodo (date) e numeri/percentuali/importi non cambiano mai
  - FORMAZIONE/CERTIFICAZIONI non vengono mai toccate
  - se una qualunque fase fallisce (LLM, template mancante, ecc.) non
    viene generato/allegato nulla — mai un CV non verificato

Il modulo genera solo il .docx personalizzato (semi-automatico): la
conversione automatica in PDF via LibreOffice è stata provata e scartata
(vedi nota sopra la funzione genera_cv_personalizzato) perché rompeva il
layout a due colonne del documento. L'ultimo passaggio — apertura in
Word ed esportazione in PDF — resta manuale per garantire la fedeltà
grafica del CV inviato a un vero datore di lavoro.

Il documento sorgente è cv_template/Giovanni Ghigliotti CV___2026.docx.
Gli indici dei paragrafi bullet modificabili sono stati verificati
manualmente il 2026-07-16: l'ordine dei paragrafi nell'XML del docx non
coincide con l'ordine di lettura visivo (il layout a due colonne interlaccia
paragrafi della sidebar sinistra con quelli della colonna destra), quindi
non è affidabile derivarli a runtime da un semplice range di indici.
"""
import os
import json
import logging
import tempfile
import hashlib

import docx
from dotenv import load_dotenv
import llm_utils

try:
    import anthropic
    _ANTHROPIC_SDK_AVAILABLE = True
except ImportError:
    _ANTHROPIC_SDK_AVAILABLE = False

load_dotenv(override=True)

_BASE_DIR = os.path.dirname(os.path.abspath(__file__))
CV_DOCX_TEMPLATE = os.path.join(_BASE_DIR, "cv_template", "Giovanni Ghigliotti CV___2026.docx")
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY", "").strip()
MATCH_LLM_MODEL = "claude-sonnet-5"

# Indici verificati manualmente (vedi docstring del modulo).
INDICE_TITLE = 15
INDICE_CITTA = 13
INDICE_INTRO = 18
CITTA_SUPPORTATE = {"Genova", "Milano", "Torino"}
INDICI_TITOLI_RUOLO = [20, 38, 53, 68, 82]
INDICI_BULLET_MODIFICABILI = [
    35, 36, 37,
    48, 49, 50, 51,
    63, 64, 65, 66,
    75, 76, 77, 78, 79, 80, 81,
    88, 89, 90, 91, 92, 93, 94, 95, 96, 97,
]
INDICI_MODIFICABILI = INDICI_TITOLI_RUOLO + INDICI_BULLET_MODIFICABILI

# Hash SHA-256 dell'intera sequenza di testo dei paragrafi del template,
# calcolato l'ultima volta che gli indici INDICE_*/INDICI_* sono stati
# verificati manualmente contro il file reale (2026-07-17). La vecchia
# validazione per "ancore" (5 coppie indice-titolo-ruolo/azienda-vicina)
# copriva solo quei 5 indici, lasciando scoperti gli altri 27 indici bullet
# e title/città/intro — la stragrande maggioranza di ciò che viene davvero
# riscritto — dando una falsa sicurezza. L'hash copre l'INTERA struttura:
# qualunque modifica al template, ovunque, lo cambia, bloccando la
# personalizzazione finché gli indici non vengono ri-verificati a mano e
# questa costante aggiornata di conseguenza (vedi _hash_template più sotto).
_TEMPLATE_HASH_ATTESO = "7d67a356e02ee12ff4fe003a23dd18b529e8e36d9b725c3c92a7e164dcdb1552"

def _hash_template(doc):
    """Hash della sequenza di testo di tutti i paragrafi, per rilevare qualunque
    modifica al template — anche una che non cambia il conteggio totale dei
    paragrafi e che quindi il vecchio controllo di sola lunghezza non vedrebbe."""
    testi = [p.text for p in doc.paragraphs]
    return hashlib.sha256("\n".join(testi).encode("utf-8")).hexdigest()

_client = None

def _get_client():
    global _client
    if _client is None and _ANTHROPIC_SDK_AVAILABLE and ANTHROPIC_API_KEY:
        # timeout ridotto da 180s a 90s per limitare il caso peggiore di due chiamate
        # sequenziali (proposta+verifica). max_retries lasciato al default SDK (2, cioè
        # 3 tentativi): un giro precedente lo aveva ridotto a 1, ma quel limite di tempo
        # non era comunque garantito (il budget CV_PERSONALIZZAZIONE_BUDGET_SECONDI
        # controlla solo se avviare il job successivo, non una chiamata già in corso),
        # mentre ridurre i retry peggiora misurabilmente la resilienza contro gli errori
        # 529 Overloaded osservati realmente durante test su questo stesso progetto.
        _client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY, timeout=90.0)
    return _client


# ==========================================
# UTILITY DOCX
# ==========================================
def formatta_title_spaziato(titolo: str) -> str:
    """Riproduce il pattern del title del CV: lettere separate da spazio
    dentro ogni parola, parole separate da tab — es. 'Head of Growth' diventa
    'H e a d\\tof\\tG r o w t h' seguendo esattamente il pattern originale
    ('D i g i t a l\\tS a l e s\\t&\\tM a r k e t i n g\\tM a n a g e r')."""
    parole = titolo.split()
    return "\t".join(" ".join(list(p)) for p in parole)


def _sostituisci_testo_paragrafo(paragraph, nuovo_testo):
    """Sostituisce il testo di un paragrafo mantenendo la formattazione
    (font, size, bold, colore) del primo run esistente. Se il paragrafo non
    ha run (caso anomalo per un template già compilato), il testo viene
    comunque inserito ma con lo stile di default: la formattazione originale
    NON è preservabile in quel caso, e la funzione lo segnala esplicitamente
    invece di fallire in silenzio."""
    if not paragraph.runs:
        logging.warning(
            f"_sostituisci_testo_paragrafo: paragrafo senza run esistenti, "
            f"formattazione originale non preservabile per il testo: {nuovo_testo[:60]!r}"
        )
        paragraph.add_run(nuovo_testo)
        return
    primo = paragraph.runs[0]
    bold, italic, underline = primo.bold, primo.italic, primo.underline
    font_name = primo.font.name
    font_size = primo.font.size
    font_color = None
    try:
        if primo.font.color and primo.font.color.type is not None:
            font_color = primo.font.color.rgb
    except Exception:
        pass
    for extra in paragraph.runs[1:]:
        extra.text = ""
    primo.text = nuovo_testo
    primo.bold, primo.italic, primo.underline = bold, italic, underline
    if font_name:
        primo.font.name = font_name
    if font_size:
        primo.font.size = font_size
    if font_color:
        primo.font.color.rgb = font_color


# ==========================================
# STEP 1 — PROPOSTA MODIFICHE (Claude)
# ==========================================
_SYSTEM_PROPOSTA = """Sei un copywriter esperto di CV: il tuo compito è riformulare i contenuti di un CV per massimizzare l'allineamento con un annuncio di lavoro specifico, puntando a un match del 100%. L'utente farà una revisione personale di ogni proposta prima di usarla per candidarsi, quindi punta a un risultato incisivo e ben calibrato sull'annuncio, non a un cambiamento minimo.

Puoi riformulare liberamente: enfasi, linguaggio, ordine dei concetti, ed etichette di ruolo (i titoli di ciascuna esperienza) per usare terminologia più riconoscibile e allineata all'annuncio — restando comunque coerente con le responsabilità realmente descritte in quell'esperienza.

L'UNICO VINCOLO RIGIDO: non introdurre strumenti, tecnologie, piattaforme o certificazioni nominate che non compaiono già da nessuna parte nel CV originale (es. se l'annuncio chiede Salesforce e il CV non lo cita in nessuna esperienza, non aggiungerlo — ometti semplicemente quel punto specifico, il resto della riformulazione procede normalmente).

Altre regole pratiche:
- Non toccare numeri, percentuali, importi o date: vanno riportati identici.
- Non toccare il nome dell'azienda in nessuna esperienza.
- Per il paragrafo PROFILO PROFESSIONALE: mantieni lo stesso framework generale (titolo professionale, track record, azioni chiave, chiusura su cosa cerca), riformulando enfasi e linguaggio per allinearli all'annuncio.
- Per i TITOLI DI RUOLO: puoi usare un'etichetta diversa per la stessa posizione realmente ricoperta (es. se l'annuncio cerca "Head of Growth" e un'esperienza reale copre effettivamente quelle responsabilità, puoi rietichettarla), ma non cambiare la sostanza del ruolo svolto.

Per ogni modifica che proponi, spiega in "fonte" da quale punto del CV originale deriva (per la verifica indipendente che segue)."""

def _proponi_modifiche_cv(job_title, job_text, intro_originale, bullet_originali):
    client = _get_client()
    if client is None:
        return None

    bullet_elenco = "\n".join(
        f"[{idx}] {'(TITOLO DI RUOLO) ' if idx in INDICI_TITOLI_RUOLO else ''}{testo}"
        for idx, testo in bullet_originali.items()
    )
    user_content = f"""TITOLO DELL'ANNUNCIO: {job_title}

TESTO INTEGRALE DELL'ANNUNCIO:
{job_text[:8000]}

PARAGRAFO PROFILO PROFESSIONALE ORIGINALE:
{intro_originale}

TITOLI DI RUOLO E BULLET DELLE ESPERIENZE ORIGINALI (indice: testo):
{bullet_elenco}

Proponi le riformulazioni che avvicinano il CV al linguaggio di questo annuncio, puntando al massimo allineamento possibile, seguendo tutte le regole del tuo ruolo."""

    try:
        with client.messages.stream(
            model=MATCH_LLM_MODEL,
            max_tokens=16000,
            system=_SYSTEM_PROPOSTA,
            messages=[{"role": "user", "content": user_content}],
            output_config={
                "effort": "medium",
                "format": {
                    "type": "json_schema",
                    "schema": {
                        "type": "object",
                        "properties": {
                            "intro_riformulato": {"type": "string"},
                            "modifiche_bullet": {
                                "type": "array",
                                "items": {
                                    "type": "object",
                                    "properties": {
                                        "indice": {"type": "integer"},
                                        "nuovo_testo": {"type": "string"},
                                        "fonte": {"type": "string"},
                                    },
                                    "required": ["indice", "nuovo_testo", "fonte"],
                                    "additionalProperties": False,
                                },
                            },
                        },
                        "required": ["intro_riformulato", "modifiche_bullet"],
                        "additionalProperties": False,
                    },
                },
            },
        ) as stream:
            response = stream.get_final_message()
        testo = llm_utils.estrai_testo_risposta(response)
        return json.loads(testo)
    except Exception as e:
        logging.warning(f"Proposta modifiche CV fallita: {type(e).__name__}: {e}")
        return None


# ==========================================
# STEP 2 — VERIFICA ANTI-FABBRICAZIONE (Claude, ruolo separato)
# ==========================================
_SYSTEM_VERIFICA = """Sei un revisore indipendente. L'utente ha scelto esplicitamente di dare ampia libertà di riformulazione (enfasi, linguaggio, etichette di ruolo) e farà lui stesso una revisione finale prima di ogni candidatura — quindi APPROVA per default le riformulazioni stilistiche, di enfasi o le rietichettature di ruolo, anche se non sono una parafrasi letterale del CV originale.

L'UNICO MOTIVO PER RESPINGERE (approvato=false): il testo proposto nomina esplicitamente uno strumento, tecnologia, piattaforma o certificazione specifica (es. "Salesforce", "SAP", "AWS certificato", ecc.) che non compare da NESSUNA parte nel CV originale integrale fornito. In questo caso, e solo in questo caso, respingi e nel testo_finale riporta il testo ORIGINALE invariato.

Respingi anche se il testo proposto cambia numeri, percentuali, importi, date o il nome di un'azienda rispetto all'originale — questi restano sempre invariati.

Per tutto il resto (riformulazioni di enfasi, linguaggio, struttura, etichette di ruolo che restano coerenti con le responsabilità reali dell'esperienza) approva. Non correggere né migliorare il testo proposto: approvi o respingi così com'è."""

def _verifica_modifiche_cv(cv_testo_completo, proposta):
    client = _get_client()
    if client is None:
        return None

    modifiche_elenco = "\n".join(
        f"[{m['indice']}] PROPOSTO: {m['nuovo_testo']}\n     FONTE DICHIARATA: {m['fonte']}"
        for m in proposta.get("modifiche_bullet", [])
    )
    user_content = f"""CV ORIGINALE INTEGRALE (verità di riferimento):
{cv_testo_completo[:8000]}

PROFILO PROFESSIONALE PROPOSTO (da verificare):
{proposta.get('intro_riformulato', '')}

BULLET PROPOSTI (da verificare uno per uno):
{modifiche_elenco}

Verifica ogni elemento secondo le tue regole."""

    try:
        with client.messages.stream(
            model=MATCH_LLM_MODEL,
            max_tokens=16000,
            system=_SYSTEM_VERIFICA,
            messages=[{"role": "user", "content": user_content}],
            output_config={
                "effort": "medium",
                "format": {
                    "type": "json_schema",
                    "schema": {
                        "type": "object",
                        "properties": {
                            "intro": {
                                "type": "object",
                                "properties": {
                                    "approvato": {"type": "boolean"},
                                    "testo_finale": {"type": "string"},
                                },
                                "required": ["approvato", "testo_finale"],
                                "additionalProperties": False,
                            },
                            "bullet": {
                                "type": "array",
                                "items": {
                                    "type": "object",
                                    "properties": {
                                        "indice": {"type": "integer"},
                                        "approvato": {"type": "boolean"},
                                        "testo_finale": {"type": "string"},
                                    },
                                    "required": ["indice", "approvato", "testo_finale"],
                                    "additionalProperties": False,
                                },
                            },
                        },
                        "required": ["intro", "bullet"],
                        "additionalProperties": False,
                    },
                },
            },
        ) as stream:
            response = stream.get_final_message()
        testo = llm_utils.estrai_testo_risposta(response)
        return json.loads(testo)
    except Exception as e:
        logging.warning(f"Verifica modifiche CV fallita: {type(e).__name__}: {e}")
        return None


# ==========================================
# ORCHESTRAZIONE
# ==========================================
# NOTA: la conversione automatica .docx -> .pdf via LibreOffice è stata
# provata e scartata (verificato in CI il 2026-07-16): il motore di
# rendering di LibreOffice non riproduce fedelmente il layout a due
# colonne/frame del documento sorgente — risultato un PDF di 4 pagine con
# contenuti duplicati invece dell'originale a 1 pagina. Il flusso è quindi
# semi-automatico: questo modulo genera e verifica il .docx, l'ultimo
# passaggio (apertura in Word ed esportazione PDF) resta manuale per
# garantire la fedeltà grafica del CV inviato a un vero datore di lavoro.
def genera_cv_personalizzato(job_title: str, job_text: str, job_city: str = None, output_dir: str = None):
    """Genera un CV .docx personalizzato per un annuncio specifico.
    Ritorna {"docx_path": ..., "riepilogo": [righe leggibili]} oppure None
    se una qualunque fase fallisce — in tal caso nessun file viene
    generato/allegato, mai un CV non verificato."""
    if not os.path.exists(CV_DOCX_TEMPLATE):
        logging.warning(f"Template CV .docx non trovato: {CV_DOCX_TEMPLATE}")
        return None
    if not job_title or not job_text or len(job_text.strip()) < 50:
        logging.warning("Testo annuncio insufficiente per la personalizzazione del CV.")
        return None

    try:
        doc = docx.Document(CV_DOCX_TEMPLATE)
    except Exception as e:
        logging.error(f"Errore apertura template CV: {e}")
        return None

    if len(doc.paragraphs) <= max(INDICI_MODIFICABILI + [INDICE_TITLE, INDICE_CITTA, INDICE_INTRO]):
        logging.error("Il template CV .docx non ha la struttura attesa (paragrafi mancanti) — personalizzazione annullata.")
        return None

    hash_attuale = _hash_template(doc)
    if hash_attuale != _TEMPLATE_HASH_ATTESO:
        logging.error(
            f"Il template CV .docx risulta modificato rispetto all'ultima verifica manuale degli "
            f"indici (hash atteso {_TEMPLATE_HASH_ATTESO[:12]}…, trovato {hash_attuale[:12]}…) — "
            f"personalizzazione annullata per non rischiare di scrivere nel paragrafo sbagliato. "
            f"Se il template è stato modificato intenzionalmente, ri-verificare manualmente tutti "
            f"gli indici INDICE_*/INDICI_* e aggiornare _TEMPLATE_HASH_ATTESO."
        )
        return None

    paragrafi_originali = {i: doc.paragraphs[i].text.strip() for i in INDICI_MODIFICABILI}
    intro_originale = doc.paragraphs[INDICE_INTRO].text.strip()
    cv_testo_completo = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    proposta = _proponi_modifiche_cv(job_title, job_text, intro_originale, paragrafi_originali)
    if proposta is None:
        return None

    verificata = _verifica_modifiche_cv(cv_testo_completo, proposta)
    if verificata is None:
        return None

    riepilogo = []

    # Title: sostituzione deterministica, sempre applicata (non richiede verifica LLM,
    # è il titolo dell'annuncio stesso, non un'affermazione sul candidato)
    _sostituisci_testo_paragrafo(doc.paragraphs[INDICE_TITLE], formatta_title_spaziato(job_title))
    riepilogo.append(f'Title aggiornato a: "{job_title}"')

    # Città nella sidebar: dinamica solo per le città seguite, deterministica
    # (non è un'affermazione su un'esperienza, non richiede verifica LLM)
    if job_city and job_city in CITTA_SUPPORTATE:
        citta_originale = doc.paragraphs[INDICE_CITTA].text.strip()
        if job_city != citta_originale:
            _sostituisci_testo_paragrafo(doc.paragraphs[INDICE_CITTA], job_city)
            riepilogo.append(f'Città aggiornata a: "{job_city}"')

    modifiche_llm_applicate = 0

    # Il testo effettivamente scritto nel CV è SEMPRE quello della PROPOSTA (stage 1,
    # passato dal controllo "fonte"-grounded anti-fabbricazione), mai il testo_finale
    # della VERIFICA (stage 2): la verifica può solo approvare/respingere un indice,
    # non riscriverne il contenuto — altrimenti una fabbricazione introdotta nello
    # stage di verifica stesso aggirerebbe interamente il controllo anti-fabbricazione
    # che è il vincolo rigido dichiarato di questo modulo.
    intro_info = verificata.get("intro", {})
    if intro_info.get("approvato"):
        nuovo_intro = (proposta.get("intro_riformulato") or "").strip()
        if nuovo_intro and nuovo_intro != intro_originale:
            _sostituisci_testo_paragrafo(doc.paragraphs[INDICE_INTRO], nuovo_intro)
            riepilogo.append("Profilo professionale riformulato per l'annuncio")
            modifiche_llm_applicate += 1

    # Indici e testi realmente proposti nello stage 1: lo stage 2 (verifica) deve
    # poter solo approvare/respingere quelli, mai introdurne di nuovi né riscriverne
    # il contenuto (vedi commento sopra).
    testi_proposti = {m.get("indice"): m.get("nuovo_testo", "") for m in proposta.get("modifiche_bullet", [])}

    for voce in verificata.get("bullet", []):
        idx = voce.get("indice")
        if idx not in INDICI_MODIFICABILI or idx not in testi_proposti or not voce.get("approvato"):
            continue
        nuovo = (testi_proposti[idx] or "").strip()
        originale = paragrafi_originali.get(idx, "")
        if not nuovo or nuovo == originale:
            continue
        _sostituisci_testo_paragrafo(doc.paragraphs[idx], nuovo)
        etichetta = "Titolo di ruolo" if idx in INDICI_TITOLI_RUOLO else "Bullet"
        riepilogo.append(f'{etichetta} riformulato: "{originale[:70]}..." -> "{nuovo[:70]}..."')
        modifiche_llm_applicate += 1

    if modifiche_llm_applicate == 0:
        logging.info("Nessuna modifica di contenuto verificata per questo annuncio — CV personalizzato non generato.")
        return None

    if output_dir is None:
        output_dir = tempfile.mkdtemp(prefix="cv_personalizzato_")
    else:
        os.makedirs(output_dir, exist_ok=True)

    docx_out = os.path.join(output_dir, "cv_personalizzato.docx")
    try:
        doc.save(docx_out)
    except Exception as e:
        logging.error(f"Errore salvataggio docx personalizzato: {e}")
        return None

    return {"docx_path": docx_out, "riepilogo": riepilogo}


def genera_cv_per_offerta(job_title: str, job_link: str, job_city: str = None, output_dir: str = None, job_text: str = None):
    """Punto d'ingresso usato dallo scraper. Se il chiamante passa già
    job_text (il testo integrale scaricato durante lo scoring iniziale,
    persistito in ScrapedJob.testo_completo), lo riusa senza riscaricare
    l'annuncio: tra scraping ed email possono passare ore, nel frattempo
    l'annuncio può essere stato rimosso o modificato. Riscarica da job_link
    solo se job_text non è disponibile (fallback per compatibilità).
    Ritorna None su qualunque fallimento, senza sollevare eccezioni verso il chiamante."""
    if not job_text or not job_text.strip():
        try:
            import requests
            from bs4 import BeautifulSoup
            from scraper import _url_is_safe_to_fetch, USER_AGENT_CHROME
            if not _url_is_safe_to_fetch(job_link):
                logging.warning(f"URL scartato (non http/https o punta a un indirizzo privato/interno): {job_link}")
                return None
            headers = {"User-Agent": USER_AGENT_CHROME}
            resp = requests.get(job_link, headers=headers, timeout=10)
            if resp.status_code != 200:
                logging.warning(f"Impossibile riscaricare l'annuncio per la personalizzazione CV ({job_link}): HTTP {resp.status_code}")
                return None
            soup = BeautifulSoup(resp.text, "html.parser")
            job_text = soup.get_text(" ", strip=True)
        except Exception as e:
            logging.warning(f"Errore riscaricamento annuncio per personalizzazione CV: {e}")
            return None

    try:
        return genera_cv_personalizzato(job_title, job_text, job_city=job_city, output_dir=output_dir)
    except Exception as e:
        logging.error(f"Errore imprevisto nella generazione del CV personalizzato: {e}")
        return None
